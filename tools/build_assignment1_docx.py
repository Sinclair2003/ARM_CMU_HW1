from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("submission/assignment1_report_draft.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def format_table(table, header_fill="E8EEF5"):
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        hdr[idx].text = text
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = str(text)
    set_table_width(table, widths)
    format_table(table)
    doc.add_paragraph()
    return table


def add_code_block(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9.5)


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(10)

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10 if style_name != "Heading 1" else 14)
        style.paragraph_format.space_after = Pt(6)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    configure_styles(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Assignment 1: Cointegration & Pairs Trading")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Concise Report Draft")
    run.italic = True
    run.font.color.rgb = RGBColor(85, 85, 85)

    p = doc.add_paragraph()
    p.add_run("Team Information: ").bold = True
    p.add_run("[Fill in team member names and Andrew IDs]")

    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(
        "This report applies a two-stage cointegration and pairs trading workflow to two selected pairs. "
        "Daily adjusted close prices were downloaded from Yahoo Finance using the yfinance Python package. "
        "To reduce look-ahead bias, the sample is separated into a formation period and a trading period."
    )

    add_table(
        doc,
        ["Category", "Pair", "Instruments"],
        [
            ["Stock pair", "MU / WDC", "Micron Technology and Western Digital"],
            ["ETF pair", "EWA / EWC", "iShares MSCI Australia ETF and iShares MSCI Canada ETF"],
        ],
        [1.25, 1.25, 4.0],
    )

    add_table(
        doc,
        ["Period", "Dates", "Purpose"],
        [
            ["Formation period", "2021-01-04 to 2023-12-29", "Correlation, cointegration testing, hedge-ratio estimation"],
            ["Trading period", "2024-01-02 to 2026-05-12", "Out-of-sample strategy backtest"],
        ],
        [1.55, 2.0, 2.95],
    )

    doc.add_paragraph("The spread is constructed as:")
    add_code_block(doc, "spread = first asset - alpha - beta * second asset")
    doc.add_paragraph(
        "The hedge ratio beta and intercept alpha are estimated by OLS during the formation period "
        "and then fixed during the trading period."
    )

    doc.add_heading("2. Pair 1: MU / WDC", level=1)
    doc.add_heading("Economic Motivation", level=2)
    doc.add_paragraph(
        "Micron Technology (MU) and Western Digital (WDC) are both exposed to the memory and data storage cycle. "
        "Micron is strongly linked to DRAM and NAND memory, while Western Digital is exposed to NAND flash and "
        "storage devices. Both firms are affected by semiconductor demand, cloud and data-center investment, "
        "inventory cycles, and pricing pressure in memory/storage markets. This makes a long-run relationship "
        "plausible, although it must be verified empirically."
    )

    doc.add_heading("Preliminary Screening", level=2)
    doc.add_paragraph(
        "During the formation period, the price-level correlation between MU and WDC was 0.7997. "
        "This indicates meaningful co-movement, but correlation alone is not sufficient because two "
        "non-stationary assets can appear correlated without having a stable mean-reverting spread."
    )

    doc.add_heading("Cointegration Results", level=2)
    add_table(
        doc,
        ["Metric", "Result"],
        [
            ["Test statistic", "-2.1042"],
            ["p-value", "0.4742"],
            ["OLS alpha", "31.9935"],
            ["OLS beta", "0.9943"],
        ],
        [3.25, 3.25],
    )
    doc.add_paragraph(
        "The null hypothesis is no cointegration. Since the p-value is above conventional 10%, 5%, and 1% "
        "levels, we do not reject the null hypothesis during the formation period. Therefore, the strict "
        "split-sample evidence for cointegration is weak."
    )

    doc.add_heading("Strategy Design", level=2)
    doc.add_paragraph("The trading strategy uses the fixed formation-period spread:")
    add_code_block(doc, "spread_t = MU_t - 31.9935 - 0.9943 * WDC_t")
    add_table(
        doc,
        ["Signal", "Action"],
        [
            ["z-score > 2.0", "Short spread: short MU, long beta-adjusted WDC"],
            ["z-score < -2.0", "Long spread: long MU, short beta-adjusted WDC"],
            ["abs(z-score) < 0.5", "Close position"],
        ],
        [2.2, 4.3],
    )
    doc.add_paragraph("Positions are shifted by one day before calculating returns to avoid look-ahead bias.")

    doc.add_heading("Backtesting Results", level=2)
    add_table(
        doc,
        ["Metric", "Result"],
        [
            ["Trading period", "2024-01-02 to 2026-05-12"],
            ["Total return", "39.98%"],
            ["Annualized Sharpe ratio", "0.6299"],
            ["Maximum drawdown", "-20.04%"],
        ],
        [3.25, 3.25],
    )
    doc.add_paragraph(
        "The strategy generated a positive total return, but the Sharpe ratio is moderate and the drawdown is material. "
        "Because the formation-period cointegration test did not reject the null of no cointegration, this result should "
        "be interpreted as exploratory rather than strong evidence of a persistent equilibrium relationship."
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("3. Pair 2: EWA / EWC", level=1)
    doc.add_heading("Economic Motivation", level=2)
    doc.add_paragraph(
        "EWA tracks Australian equities and EWC tracks Canadian equities. Australia and Canada are both developed, "
        "commodity-sensitive markets. Their equity markets are influenced by global growth, resource demand, energy "
        "and materials prices, exchange-rate cycles, and global risk appetite. These shared macro drivers make a "
        "long-run relationship plausible."
    )

    doc.add_heading("Preliminary Screening", level=2)
    doc.add_paragraph(
        "During the formation period, the price-level correlation between EWA and EWC was 0.7214. "
        "This shows moderate-to-high co-movement, but correlation alone does not prove that the spread is stationary."
    )

    doc.add_heading("Cointegration Results", level=2)
    add_table(
        doc,
        ["Metric", "Result"],
        [
            ["Test statistic", "-2.6586"],
            ["p-value", "0.2148"],
            ["OLS alpha", "7.4876"],
            ["OLS beta", "0.3986"],
        ],
        [3.25, 3.25],
    )
    doc.add_paragraph(
        "The null hypothesis is no cointegration. The p-value is above conventional significance thresholds, "
        "so we do not reject the null hypothesis during the formation period. The formation-period evidence for "
        "cointegration is therefore weak."
    )

    doc.add_heading("Strategy Design", level=2)
    doc.add_paragraph("The trading strategy uses the fixed formation-period spread:")
    add_code_block(doc, "spread_t = EWA_t - 7.4876 - 0.3986 * EWC_t")
    add_table(
        doc,
        ["Signal", "Action"],
        [
            ["z-score > 2.0", "Short spread: short EWA, long beta-adjusted EWC"],
            ["z-score < -2.0", "Long spread: long EWA, short beta-adjusted EWC"],
            ["abs(z-score) < 0.5", "Close position"],
        ],
        [2.2, 4.3],
    )
    doc.add_paragraph("Positions are shifted by one day before calculating returns to avoid look-ahead bias.")

    doc.add_heading("Backtesting Results", level=2)
    add_table(
        doc,
        ["Metric", "Result"],
        [
            ["Trading period", "2024-01-02 to 2026-05-12"],
            ["Total return", "15.40%"],
            ["Annualized Sharpe ratio", "0.6448"],
            ["Maximum drawdown", "-10.06%"],
        ],
        [3.25, 3.25],
    )
    doc.add_paragraph(
        "The strategy produced a positive out-of-sample return with a moderate Sharpe ratio and lower drawdown than "
        "the stock pair. However, because the formation-period cointegration evidence is weak, the result should be "
        "interpreted cautiously."
    )

    doc.add_heading("4. Summary of Results", level=1)
    add_table(
        doc,
        ["Pair", "Corr.", "Coint stat", "p-value", "Beta", "Total return", "Sharpe", "Max DD"],
        [
            ["MU / WDC", "0.7997", "-2.1042", "0.4742", "0.9943", "39.98%", "0.6299", "-20.04%"],
            ["EWA / EWC", "0.7214", "-2.6586", "0.2148", "0.3986", "15.40%", "0.6448", "-10.06%"],
        ],
        [1.0, 0.65, 0.9, 0.75, 0.65, 0.95, 0.75, 0.85],
    )

    doc.add_heading("5. Interpretation and Limitations", level=1)
    doc.add_paragraph(
        "The results illustrate the difference between economic intuition, correlation, cointegration, and trading "
        "performance. Both pairs have plausible economic relationships and positive out-of-sample backtest returns. "
        "However, under the strict split-sample design, neither pair passes the Engle-Granger cointegration test during "
        "the formation period. This means the statistical evidence for a stable long-run equilibrium is weak."
    )
    doc.add_paragraph(
        "The positive backtest performance may reflect short-term mean reversion, favorable market conditions during "
        "the trading period, or parameter choices rather than a robust cointegration relationship. The strategy should "
        "therefore be treated as an empirical experiment rather than a fully validated trading model."
    )
    add_bullets(
        doc,
        [
            "Transaction costs, bid-ask spreads, slippage, short-sale constraints, and borrowing costs are ignored.",
            "The hedge ratio is estimated once and kept fixed, although relationships can change over time.",
            "Cointegration relationships may break during regime shifts, earnings shocks, commodity cycles, or macroeconomic changes.",
            "The strategy uses daily close data and does not model intraday execution.",
            "The sample period is limited, so results may not generalize to future markets.",
        ],
    )

    doc.add_heading("Appendix: Python Code", level=1)
    doc.add_paragraph("The working assignment notebook is provided in:")
    add_code_block(doc, "notebook/assignment1_pairs_trading_report.ipynb")
    doc.add_paragraph("The equivalent Python script is provided in:")
    add_code_block(doc, "notebook/assignment1_pairs_trading_report.py")
    doc.add_paragraph(
        "The code downloads data from Yahoo Finance, computes correlations, estimates OLS hedge ratios, runs "
        "Engle-Granger cointegration tests, constructs spreads and rolling z-scores, backtests the trading strategy, "
        "and reports total return, Sharpe ratio, and maximum drawdown."
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(f"Wrote {OUT.resolve()}")
